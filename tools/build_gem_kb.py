from pathlib import Path
import re
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "raw_sources"
RAW.mkdir(exist_ok=True)

SOURCES = {
    "arabic": [
        (1, "كفايات اللغة العربية ومعاييرها للصف1. .docx"),
        (2, "كفايات اللغة العربية ومعاييرها للصف 2..docx"),
        (3, "كفايات اللغة العربية ومعاييرها للصف 3. .docx"),
        (4, "كفايات اللغة العربية ومعاييرها للصف 4..docx"),
    ],
    "math": [
        (1, "كفايات الرياضيات ومعاييرها للصف1.docx"),
        (2, "كفايات الرياضيات ومعاييرها للصف 2..docx"),
        (3, "كفايات الرياضيات ومعاييرها للصف 3.docx"),
        (4, "كفايات الرياضيات ومعاييرها للصف 4..docx"),
    ],
    "islamic": [
        (1, "كفايات التربية الاسلامية ومعاييرها للصف 1. .docx"),
        (2, "كفايات التربية الاسلامية ومعاييرها للصف 2. ا.docx"),
        (3, "كفايات التربية الاسلامية ومعاييرها للصف 3. docx.docx"),
        (4, "كفايات التربية الاسلامية ومعاييرها للصف 4. .docx"),
    ],
    "nurturing": [
        (1, "الكفايات والمعايبر والمؤشرات والقيم علوم وطنية وحياتية  اول (1.docx"),
        (2, "الكفايات والمعايبر والمؤشرات والقيم علوم وطنية وحياتية ثاني (1).docx"),
        (3, "كفايات التنشئة ومعاييرها للصف 3. .docx"),
        (4, "كفايات التنشئة ومعاييرها للصف 4. .docx"),
    ],
}


def norm(s):
    s = s.replace("\u00a0", " ").replace("\u200f", "").replace("\u200e", "")
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def extract_docx(path):
    doc = Document(path)
    blocks = []
    for p in doc.paragraphs:
        t = norm(p.text)
        if t:
            blocks.append(t)
    for ti, table in enumerate(doc.tables, 1):
        blocks.append(f"[TABLE {ti}]")
        for row in table.rows:
            cells = [norm(c.text).replace("\n", " / ") for c in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return blocks


def get_blocks(subject, grade, filename):
    docx = ROOT / filename
    raw = RAW / f"{subject}_grade_{grade}.txt"
    if docx.exists():
        blocks = extract_docx(docx)
        raw.write_text("\n".join(blocks) + "\n", encoding="utf-8")
        return blocks, "DOCX"
    if raw.exists():
        return raw.read_text(encoding="utf-8").splitlines(), "RAW_TEXT"
    return [f"SOURCE FILE MISSING: `{filename}`"], "MISSING"


def build_subject(subject, items):
    out = [
        f"# GEM KNOWLEDGE BASE — {subject.upper()} — GRADES 1–4",
        "",
        "STATUS: GENERATED FROM USER-PROVIDED REFERENCE MATERIAL.",
        "SOURCE CLASS: USER-PROVIDED REFERENCE; NOT CLAIMED OFFICIAL-VERIFIED.",
        "RULE: PRESERVE SOURCE TEXT; NORMALIZATION MUST NOT replace the raw source.",
        "TRACE: SOURCE → COMPETENCY → LEARNING OUTCOME → OBSERVABLE PERFORMANCE → ACTIVITY → EVIDENCE → ASSESSMENT.",
        "",
    ]
    for grade, filename in items:
        blocks, origin = get_blocks(subject, grade, filename)
        out += [f"## الصف {grade}", f"SOURCE: `{filename}`", f"EXTRACTION ORIGIN: {origin}", "", "### ORIGINAL EXTRACTED TEXT", ""]
        out += blocks
        out += ["", "---", ""]
    return "\n".join(out) + "\n"

for subject, items in SOURCES.items():
    target = ROOT / f"GEM_KB_{subject.upper()}_GRADES_1-4.md"
    target.write_text(build_subject(subject, items), encoding="utf-8")

status = ROOT / "BASELINE_EXTRACTION_STATUS.md"
status.write_text(
    "# BASELINE EXTRACTION STATUS\n\n"
    "Generated automatically from repository DOCX sources when available, otherwise from previously extracted raw text.\n\n"
    "All extracted text is retained under `raw_sources/`; subject KBs retain extracted text grouped by grade.\n"
    "Missing source files are recorded explicitly; the builder never fabricates competency content.\n"
    "The nurturing grade 1–2 curriculum naming transition is preserved as a source mapping and is not presented as independently verified equivalence.\n",
    encoding="utf-8",
)
